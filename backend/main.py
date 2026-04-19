import sys, os
os.environ["PYTHONIOENCODING"] = "utf-8"
os.environ["PYTHONUTF8"] = "1"
for _s in (sys.stdout, sys.stderr):
    if _s and hasattr(_s, "reconfigure"):
        try: _s.reconfigure(encoding="utf-8", errors="replace")
        except Exception: pass

from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import pdfplumber
import docx
import json
import io
import os
from pathlib import Path

from database import get_db, init_db
from models import CV, Job, Evaluation, User
from schemas import CVResponse, JobCreate, JobResponse, EvaluateRequest, EvaluationResponse
from claude_service import extract_cv_data, evaluate_cv, evaluate_atc, hr_review, search_jobs_chile, improve_cv
from auth import hash_password, verify_password, create_token, get_current_user, require_auth, ADMIN_EMAIL

app = FastAPI(title="HR CV Evaluator API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

FRONTEND_DIR = Path(__file__).parent / "frontend"
if not FRONTEND_DIR.exists():
    FRONTEND_DIR = Path(__file__).parent.parent / "frontend"
if FRONTEND_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(FRONTEND_DIR)), name="static")


@app.on_event("startup")
async def startup():
    await init_db()


@app.get("/")
async def root():
    landing = FRONTEND_DIR / "landing.html"
    if landing.exists():
        return FileResponse(str(landing))
    return {"message": "HR CV Evaluator API", "docs": "/docs"}


# --- Auth ---

@app.post("/api/auth/register", status_code=201)
async def register(body: dict, db: AsyncSession = Depends(get_db)):
    email = body.get("email", "").lower().strip()
    password = body.get("password", "")
    name = body.get("name", "")
    if not email or not password or not name:
        raise HTTPException(status_code=400, detail="Nombre, email y contraseña son requeridos")
    if len(password) < 6:
        raise HTTPException(status_code=400, detail="La contraseña debe tener al menos 6 caracteres")
    existing = await db.execute(select(User).where(User.email == email))
    if existing.scalar_one_or_none():
        raise HTTPException(status_code=409, detail="Este email ya está registrado")
    user = User(email=email, password_hash=hash_password(password), name=name, is_admin=(email == ADMIN_EMAIL))
    db.add(user)
    await db.commit()
    await db.refresh(user)
    return {"token": create_token(user.id), "name": user.name, "email": user.email, "is_admin": user.is_admin}


@app.post("/api/auth/login")
async def login(body: dict, db: AsyncSession = Depends(get_db)):
    email = body.get("email", "").lower().strip()
    password = body.get("password", "")
    result = await db.execute(select(User).where(User.email == email))
    user = result.scalar_one_or_none()
    if not user or not verify_password(password, user.password_hash):
        raise HTTPException(status_code=401, detail="Email o contraseña incorrectos")
    return {"token": create_token(user.id), "name": user.name, "email": user.email, "is_admin": user.is_admin}


@app.get("/api/auth/me")
async def me(user: User = Depends(require_auth)):
    return {"id": user.id, "name": user.name, "email": user.email, "is_admin": user.is_admin}


@app.get("/app")
async def app_page():
    index = FRONTEND_DIR / "index.html"
    if index.exists():
        return FileResponse(str(index))
    return {"message": "App not found"}


def extract_text(file: UploadFile, content: bytes) -> str:
    name = file.filename.lower()
    if name.endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    if name.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    if name.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")
    raise HTTPException(status_code=400, detail="Formato no soportado. Usa PDF, DOCX o TXT.")


# --- CVs ---

@app.post("/api/cvs/upload", status_code=201)
async def upload_cv(file: UploadFile = File(...), db: AsyncSession = Depends(get_db), user: User = Depends(require_auth)):
    content = await file.read()
    text = extract_text(file, content)
    if not text.strip():
        raise HTTPException(status_code=422, detail="No se pudo extraer texto del archivo.")
    cv = CV(filename=file.filename, content=text, user_id=user.id)
    db.add(cv)
    await db.commit()
    await db.refresh(cv)

    atc = evaluate_atc(cv.content)
    hr = hr_review(cv.content)
    extracted = extract_cv_data(cv.content)

    cv.atc_cache = json.dumps(atc, ensure_ascii=False)
    cv.hr_cache = json.dumps(hr, ensure_ascii=False)
    cv.extracted_cache = json.dumps(extracted, ensure_ascii=False)
    await db.commit()

    return {
        "id": cv.id,
        "filename": cv.filename,
        "uploaded_at": cv.uploaded_at,
        "extracted": extracted,
        "atc": atc,
        "hr_review": hr,
    }


@app.get("/api/cvs", response_model=list[CVResponse])
async def list_cvs(db: AsyncSession = Depends(get_db), user: User = Depends(require_auth)):
    if user.is_admin:
        result = await db.execute(select(CV).order_by(CV.uploaded_at.desc()))
    else:
        result = await db.execute(select(CV).where(CV.user_id == user.id).order_by(CV.uploaded_at.desc()))
    return result.scalars().all()


@app.get("/api/cvs/{cv_id}", response_model=CVResponse)
async def get_cv(cv_id: int, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    return cv


# --- Jobs ---

@app.post("/api/jobs", response_model=JobResponse, status_code=201)
async def create_job(body: JobCreate, db: AsyncSession = Depends(get_db)):
    job = Job(**body.model_dump())
    db.add(job)
    await db.commit()
    await db.refresh(job)
    return job


@app.get("/api/jobs", response_model=list[JobResponse])
async def list_jobs(db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Job).order_by(Job.created_at.desc()))
    return result.scalars().all()


# --- Evaluaciones ---

@app.post("/api/evaluate", response_model=EvaluationResponse, status_code=201)
async def evaluate(body: EvaluateRequest, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, body.cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    job = await db.get(Job, body.job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Oferta no encontrada")

    extracted = extract_cv_data(cv.content)
    result = evaluate_cv(cv.content, job.title, job.description, job.requirements)

    evaluation = Evaluation(
        cv_id=cv.id,
        job_id=job.id,
        score=result["score"],
        feedback=json.dumps(result, ensure_ascii=False),
        extracted_data=json.dumps(extracted, ensure_ascii=False),
    )
    db.add(evaluation)
    await db.commit()
    await db.refresh(evaluation)

    return EvaluationResponse(
        **{c.name: getattr(evaluation, c.name) for c in evaluation.__table__.columns},
        cv_filename=cv.filename,
        job_title=job.title,
    )


@app.get("/api/evaluations", response_model=list[EvaluationResponse])
async def list_evaluations(db: AsyncSession = Depends(get_db)):
    result = await db.execute(
        select(Evaluation, CV.filename, Job.title)
        .join(CV, Evaluation.cv_id == CV.id)
        .join(Job, Evaluation.job_id == Job.id)
        .order_by(Evaluation.created_at.desc())
    )
    rows = result.all()
    return [
        EvaluationResponse(
            **{c.name: getattr(ev, c.name) for c in ev.__table__.columns},
            cv_filename=filename,
            job_title=job_title,
        )
        for ev, filename, job_title in rows
    ]


@app.get("/api/evaluations/{evaluation_id}", response_model=EvaluationResponse)
async def get_evaluation(evaluation_id: int, db: AsyncSession = Depends(get_db)):
    ev = await db.get(Evaluation, evaluation_id)
    if not ev:
        raise HTTPException(status_code=404, detail="Evaluación no encontrada")
    cv = await db.get(CV, ev.cv_id)
    job = await db.get(Job, ev.job_id)
    return EvaluationResponse(
        **{c.name: getattr(ev, c.name) for c in ev.__table__.columns},
        cv_filename=cv.filename if cv else None,
        job_title=job.title if job else None,
    )


# --- ATC ---

@app.post("/api/cvs/{cv_id}/atc")
async def atc_evaluation(cv_id: int, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    if cv.atc_cache:
        return json.loads(cv.atc_cache)
    result = evaluate_atc(cv.content)
    cv.atc_cache = json.dumps(result, ensure_ascii=False)
    await db.commit()
    return result


# --- HR Review ---

@app.post("/api/cvs/{cv_id}/hr-review")
async def cv_hr_review(cv_id: int, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    if cv.hr_cache:
        return json.loads(cv.hr_cache)
    result = hr_review(cv.content)
    cv.hr_cache = json.dumps(result, ensure_ascii=False)
    await db.commit()
    return result


# --- Búsqueda de empleos ---

@app.post("/api/cvs/{cv_id}/job-search")
async def job_search(cv_id: int, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    extracted = extract_cv_data(cv.content)
    jobs = search_jobs_chile(
        profession=extracted.get("main_profession", "profesional"),
        skills=extracted.get("skills", []),
        years_exp=extracted.get("years_experience", 0),
    )
    return {"perfil": extracted, "ofertas": jobs}


# --- Mejora y descarga CV ---

def _clean(text: str) -> str:
    return text.encode("utf-16", "surrogatepass").decode("utf-16")


@app.post("/api/cvs/{cv_id}/improve")
async def improve_cv_endpoint(cv_id: int, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    hr_data = json.loads(cv.hr_cache) if cv.hr_cache else hr_review(cv.content)
    improved_text = improve_cv(cv.content, hr_data)
    cv.improved_cache = improved_text
    await db.commit()

    mejoras = hr_data.get("mejoras_urgentes", [])
    keywords = hr_data.get("keywords_faltantes", [])
    scores = hr_data.get("puntuacion_por_seccion", {})
    secciones_bajas = [k.replace("_", " ").title() for k, v in scores.items() if v.get("score", 100) < 75]

    cambios = []
    if mejoras:
        cambios.append(f"Se aplicaron {len(mejoras)} mejoras urgentes detectadas por HR")
    if keywords:
        cambios.append(f"Se incorporaron {len(keywords)} keywords clave al texto")
    if secciones_bajas:
        cambios.append(f"Se reforzaron las secciones: {', '.join(secciones_bajas)}")
    cambios.append("Se agregó perfil profesional al inicio del CV")
    cambios.append("Se reemplazaron responsabilidades por logros con impacto medible")

    return {"cambios": cambios, "cv_id": cv_id}


@app.post("/api/cvs/{cv_id}/improve-download")
async def improve_download(cv_id: int, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    if not cv.improved_cache:
        raise HTTPException(status_code=400, detail="Genera primero el CV mejorado")

    improved_text = cv.improved_cache
    doc = docx.Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = docx.shared.Pt(11)
    for line in improved_text.split("\n"):
        line = _clean(line.strip())
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line.strip("*")).bold = True
        else:
            doc.add_paragraph(line)

    safe_name = _ascii_filename(cv.filename.rsplit(".", 1)[0])
    return _docx_response(doc, f"{safe_name}_mejorado.docx")


def _ascii_filename(name: str) -> str:
    return name.encode("ascii", "ignore").decode("ascii").replace(" ", "_")


def _docx_response(doc: docx.Document, filename: str) -> StreamingResponse:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe = _ascii_filename(filename)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe}"'},
    )


def _add_section(doc, title: str, items: list[str], bullet: str = "•"):
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(f"{bullet} {item}", style="List Bullet")


def _safe(text) -> str:
    return _clean(str(text)) if text else ""


@app.post("/api/cvs/{cv_id}/download-atc")
async def download_atc_report(cv_id: int, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    data = json.loads(cv.atc_cache) if cv.atc_cache else evaluate_atc(cv.content)

    doc = docx.Document()
    doc.styles["Normal"].font.name = "Calibri"

    title = doc.add_heading("Reporte Evaluación ATC", level=1)
    title.runs[0].font.color.rgb = docx.shared.RGBColor(0x16, 0xA3, 0x4A)

    doc.add_paragraph(f"Candidato: {_safe(cv.filename.rsplit('.', 1)[0])}")
    doc.add_paragraph(f"Fecha: {__import__('datetime').date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Score general: {data.get('score_general', 0)} / 100")
    doc.add_paragraph(f"Veredicto: {_safe(data.get('veredicto', ''))}")
    doc.add_paragraph("")
    doc.add_heading("Resumen", level=2)
    doc.add_paragraph(_safe(data.get("resumen", "")))

    doc.add_heading("Resultados por área", level=2)
    for key, val in data.get("tests", {}).items():
        label = key.replace("_", " ").title()
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(_safe(f"{val.get('score',0)}/100  ({val.get('nivel','')})  — {val.get('fundamento','')}")  )

    _add_section(doc, "Fortalezas ATC",   [_safe(x) for x in data.get("fortalezas_atc", [])], "+")
    _add_section(doc, "Areas de riesgo",  [_safe(x) for x in data.get("areas_riesgo", [])],   "!")
    _add_section(doc, "Como prepararse",  [_safe(x) for x in data.get("recomendaciones_preparacion", [])], "-")

    safe = cv.filename.rsplit(".", 1)[0].replace(" ", "_")
    return _docx_response(doc, f"{safe}_reporte_ATC.docx")


@app.post("/api/cvs/{cv_id}/download-hr")
async def download_hr_report(cv_id: int, db: AsyncSession = Depends(get_db)):
    cv = await db.get(CV, cv_id)
    if not cv:
        raise HTTPException(status_code=404, detail="CV no encontrado")
    data = json.loads(cv.hr_cache) if cv.hr_cache else hr_review(cv.content)

    doc = docx.Document()
    doc.styles["Normal"].font.name = "Calibri"

    title = doc.add_heading("Reporte Revisión Profesional HR", level=1)
    title.runs[0].font.color.rgb = docx.shared.RGBColor(0x16, 0xA3, 0x4A)

    doc.add_paragraph(f"Candidato: {_safe(cv.filename.rsplit('.', 1)[0])}")
    doc.add_paragraph(f"Fecha: {__import__('datetime').date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Score CV: {data.get('score_cv', 0)} / 100  -  Nivel: {_safe(data.get('nivel_cv', ''))}")
    doc.add_paragraph("")
    doc.add_heading("Resumen profesional", level=2)
    doc.add_paragraph(_safe(data.get("resumen_profesional", "")))

    doc.add_heading("Puntuacion por seccion", level=2)
    for key, val in data.get("puntuacion_por_seccion", {}).items():
        label = key.replace("_", " ").title()
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(_safe(f"{val.get('score', 0)}/100  - {val.get('comentario', '')}"))

    _add_section(doc, "Aspectos positivos", [_safe(x) for x in data.get("aspectos_positivos", [])], "+")
    _add_section(doc, "Mejoras urgentes",   [_safe(x) for x in data.get("mejoras_urgentes", [])],   "!")
    _add_section(doc, "Mejoras opcionales", [_safe(x) for x in data.get("mejoras_opcionales", [])], "-")

    if data.get("keywords_faltantes"):
        doc.add_heading("Keywords que deberias agregar", level=2)
        doc.add_paragraph(_safe("  ".join(data["keywords_faltantes"])))

    doc.add_heading("Estimacion salarial en Chile 2026", level=2)
    doc.add_paragraph(_safe(data.get("recomendacion_salarial", "")))

    doc.add_heading("Consejo para la entrevista", level=2)
    doc.add_paragraph(_safe(data.get("consejo_entrevista", "")))

    safe = cv.filename.rsplit(".", 1)[0].replace(" ", "_")
    return _docx_response(doc, f"{safe}_reporte_HR.docx")
